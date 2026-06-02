from io import BytesIO
from pathlib import Path
import re
from typing import Any

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph as DocxParagraph
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph as PdfParagraph
from reportlab.platypus import SimpleDocTemplate, Spacer


SECTION_TITLES = {
    "scope": "Scope",
    "objectives": "Objectives",
    "featuresToTest": "Features to Test",
    "featuresNotToTest": "Features Not to Test",
    "testStrategy": "Test Strategy",
    "functionalTesting": "Functional Testing",
    "nonFunctionalTesting": "Non Functional Testing",
    "securityTesting": "Security Testing",
    "apiTesting": "API Testing",
    "uiTesting": "UI Testing",
    "regressionTesting": "Regression Testing",
    "requirementsTraceability": "Requirements Traceability",
    "testEnvironment": "Test Environment",
    "entryCriteria": "Entry Criteria",
    "exitCriteria": "Exit Criteria",
    "assumptions": "Assumptions",
    "risks": "Risks",
    "deliverables": "Deliverables",
    "testCases": "Test Cases",
}

PROJECT_ROOT = Path(__file__).resolve().parents[3]
TEST_PLAN_TEMPLATE_PATH = PROJECT_ROOT / "ei_SW_TMP_TestPlan_GNRL.docx"

TEST_APPROACH_CONTEXTS = {
    9: {
        "name": "Source code/Automation Script verification",
        "objective": (
            "Verify source code, generated code where applicable, and automation scripts for "
            "compliance with project coding standards, static-analysis expectations, compiler "
            "warning policy, traceability, maintainability, and safety-related review needs."
        ),
        "goals": (
            "Detect coding-standard violations, unreviewed logic changes, script defects, "
            "traceability gaps, unresolved warnings, and safety-rule deviations before dynamic testing."
        ),
        "coverage": (
            "Coverage includes changed and safety-relevant source modules, generated code, build scripts, "
            "test automation scripts, static-analysis results, and review checklist closure."
        ),
        "technique": (
            "Use peer review, static analysis, compiler-warning review, checklist-based verification, "
            "automation dry runs, and traceability checks against the generated requirements baseline."
        ),
        "responsible": "Developer performs source-level verification with QA review for automation evidence and traceability.",
        "special": (
            "For automotive scope, include applicable coding rules, safety checklist evidence, and "
            "confirmation that reviewed changes do not violate safety goals."
        ),
        "steps": (
            "Baseline the source and scripts, run static and build checks, review changed logic against "
            "requirements, execute automation dry runs, resolve findings, and archive review evidence."
        ),
        "automation": "Automate static-analysis, build-warning, lint, and script smoke checks where supported by the project toolchain.",
    },
    10: {
        "name": "Unit Testing",
        "objective": (
            "Verify individual software units against allocated requirements, interfaces, boundary "
            "conditions, error handling, and expected outputs before integration."
        ),
        "goals": (
            "Detect defects in unit logic, boundary handling, invalid-input behavior, state decisions, "
            "diagnostic paths, and interface assumptions."
        ),
        "coverage": (
            "Coverage target is requirement-level unit coverage for impacted modules, including positive, "
            "negative, boundary, and safety-relevant paths as applicable."
        ),
        "technique": (
            "Use automated or manual unit tests with stubs/mocks, boundary-value and equivalence-class "
            "cases, decision-table cases, and fault-injection cases where required."
        ),
        "responsible": "Developer owns unit test creation and execution; QA reviews traceability and evidence.",
        "special": "Prioritize critical and safety-related requirements and document any unit-level coverage limitations.",
        "steps": (
            "Identify allocated requirements, prepare unit harness and stubs, execute nominal and negative "
            "cases, compare outputs with expected behavior, log defects, and update traceability."
        ),
        "automation": "Automate repeatable unit tests in the project unit-test framework and include them in regression where practical.",
    },
    11: {
        "name": "Unit Configuration Testing",
        "objective": (
            "Verify unit behavior across applicable calibration, configuration, compile-time, and "
            "variant settings that can affect requirement behavior."
        ),
        "goals": "Detect configuration-specific defects, missing defaults, invalid ranges, and variant handling gaps.",
        "coverage": (
            "Coverage includes valid, boundary, default, and invalid configuration values for impacted "
            "software units and safety-relevant parameters."
        ),
        "technique": (
            "Run parameterized unit tests and reviews across supported configuration sets using boundary "
            "analysis and negative configuration cases."
        ),
        "responsible": "Developer executes configuration-focused unit tests with QA review of coverage and evidence.",
        "special": "Keep configuration baselines synchronized with the SRS, calibration data, and integration environment.",
        "steps": (
            "Select configuration values, execute unit behavior under each set, verify outputs and error "
            "handling, compare with acceptance criteria, and archive results."
        ),
        "automation": "Automate parameterized configuration tests for changed modules and high-risk variants.",
    },
    12: {
        "name": "Integration Testing",
        "objective": (
            "Verify interactions among integrated software units, middleware, diagnostics, communication "
            "interfaces, and external software or hardware interfaces described by the source document."
        ),
        "goals": "Detect interface mismatches, integration timing defects, data-mapping errors, and cross-module state defects.",
        "coverage": (
            "Coverage includes impacted interfaces, message/signals, state transitions, diagnostics, "
            "fault paths, and cross-module requirement flows."
        ),
        "technique": (
            "Use SIL/HIL or bench integration tests, interface simulators, diagnostic tools, log review, "
            "positive and negative interface cases, and fault injection."
        ),
        "responsible": "QA test engineer executes integration tests with developer and systems support for interface analysis.",
        "special": "Synchronize interface artifacts, diagnostic data, and test bench configuration before execution.",
        "steps": (
            "Deploy the integrated build, configure simulators and tools, stimulate interface scenarios, "
            "monitor cross-module behavior, verify outputs, and log defects/evidence."
        ),
        "automation": "Automate stable interface and smoke integration cases in SIL/HIL regression where the bench supports it.",
    },
    13: {
        "name": "Integration Configuration Testing",
        "objective": (
            "Verify integrated behavior across supported interface, calibration, diagnostic, and "
            "environment configurations."
        ),
        "goals": "Detect configuration incompatibilities, invalid integration defaults, and variant-specific interface failures.",
        "coverage": (
            "Coverage includes configuration sets that affect integrated requirements, communication "
            "artifacts, diagnostic behavior, timing, and safety-related behavior."
        ),
        "technique": (
            "Execute integration scenarios across selected configuration baselines using simulators, "
            "boundary values, negative cases, and log comparison."
        ),
        "responsible": "QA test engineer executes with configuration-owner support.",
        "special": "Control configuration versions and record the exact baseline used for each result.",
        "steps": (
            "Load each approved configuration, run representative integration scenarios, verify outputs "
            "and logs, compare results, and document configuration-specific defects."
        ),
        "automation": "Automate repeatable configuration sweeps for high-risk and frequently changed integration settings.",
    },
    14: {
        "name": "Software Qualification Testing",
        "objective": (
            "Verify the complete software against the SRS-derived requirements, acceptance criteria, "
            "traceability matrix, and release-quality expectations."
        ),
        "goals": "Detect requirement non-compliance, missing end-to-end behavior, release blockers, and evidence gaps.",
        "coverage": (
            "Coverage includes all generated requirement-level test cases, critical requirements, "
            "interfaces, diagnostics, non-functional behavior, and regression impacts."
        ),
        "technique": (
            "Execute requirement-based qualification tests in the approved environment using positive, "
            "negative, boundary, diagnostic, safety, and regression scenarios."
        ),
        "responsible": "QA test engineer owns qualification execution with project lead approval of results.",
        "special": "Use controlled baselines for software, test data, tools, and requirements; formally defer any unexecuted cases.",
        "steps": (
            "Baseline test inputs, execute requirement-level cases, collect objective evidence, triage "
            "defects, update RTM status, and prepare the qualification summary."
        ),
        "automation": "Automate repeatable qualification smoke and regression cases when stable and tool-supported.",
    },
    15: {
        "name": "Sanity Testing",
        "objective": "Confirm that the delivered build is stable enough for deeper verification.",
        "goals": "Detect build, boot/startup, basic communication, and critical-path failures early.",
        "coverage": "Coverage includes essential startup, communication, major functional flow, diagnostic availability, and blocker checks.",
        "technique": "Execute a small, repeatable smoke suite on every candidate build before detailed testing.",
        "responsible": "QA test engineer executes sanity testing; developer support is used for build-blocker triage.",
        "special": "Stop detailed execution when sanity fails and communicate blocker status immediately.",
        "steps": "Install or flash the candidate build, run smoke checks, verify critical outputs/logs, report blockers, and approve or reject deeper testing.",
        "automation": "Automate sanity checks wherever the environment supports fast repeatable execution.",
    },
    16: {
        "name": "Installation Testing",
        "objective": "Verify that the software build can be deployed, flashed, installed, or configured correctly in the target test environment.",
        "goals": "Detect deployment, flashing, packaging, versioning, setup, and post-install smoke defects.",
        "coverage": "Coverage includes fresh deployment, update/reflash, version confirmation, configuration loading, and post-install smoke behavior.",
        "technique": "Follow the approved deployment procedure and run a predetermined subset of smoke and functional checks after deployment.",
        "responsible": "QA test engineer executes deployment verification with build/release support as needed.",
        "special": "For embedded targets without a user installer, treat flashing, calibration loading, and bench setup as the installation workflow.",
        "steps": "Prepare the target, deploy the build, verify version/configuration, execute smoke checks, and record deployment evidence.",
        "automation": "Automate deployment verification and version checks where supported by tooling.",
    },
    17: {
        "name": "Functional Testing",
        "objective": "Verify SRS functional requirements, feature behavior, inputs, outputs, state changes, and expected results.",
        "goals": "Detect functional non-compliance, incorrect state behavior, missing outputs, and invalid requirement handling.",
        "coverage": "Coverage includes generated functional test cases, requirement IDs, positive/negative paths, boundary conditions, and acceptance criteria.",
        "technique": "Execute requirement-based functional tests with documented preconditions, test data, steps, and expected results.",
        "responsible": "QA test engineer executes functional testing with developer support for defect analysis.",
        "special": "Maintain bidirectional traceability between requirements, test cases, results, and defects.",
        "steps": "Select requirements, configure preconditions, execute test steps, observe outputs, compare expected results, and update RTM/defects.",
        "automation": "Automate stable functional cases that are repeatable and valuable for regression.",
    },
    18: {
        "name": "Performance Testing",
        "objective": "Verify timing, response, throughput, latency, and resource-related behavior identified or implied by the source requirements.",
        "goals": "Detect timing violations, slow responses, resource bottlenecks, and performance regressions.",
        "coverage": "Coverage includes non-functional timing and latency requirements, communication timing, diagnostic timing, and high-risk performance paths.",
        "technique": "Measure behavior with representative workloads, logs, timestamps, profiling tools, and acceptance thresholds.",
        "responsible": "QA test engineer executes performance tests with systems/tool support for measurement setup.",
        "special": "Use calibrated tools and stable test conditions so measurements are repeatable.",
        "steps": "Prepare workload and measurement tools, execute performance scenarios, capture metrics, compare thresholds, and document deviations.",
        "automation": "Automate repeatable performance measurements where the environment and tools provide stable data.",
    },
    19: {
        "name": "Load Testing",
        "objective": "Verify that the software continues to operate correctly under expected and elevated operational workloads.",
        "goals": "Detect failures under sustained traffic, repeated state changes, high message rates, or heavy diagnostic/processing activity.",
        "coverage": "Coverage includes representative workload ranges, sustained operation, communication loads, and recovery after load removal.",
        "technique": "Apply controlled workloads with simulators or bench tools while monitoring outputs, timing, faults, and logs.",
        "responsible": "QA test engineer executes load testing with bench/tool support.",
        "special": "Confirm load levels and acceptance limits with stakeholders when they are not explicit in the source document.",
        "steps": "Baseline nominal behavior, increase workload to selected levels, monitor behavior, verify recovery, and record metrics/defects.",
        "automation": "Automate workload generation and log collection where possible.",
    },
    20: {
        "name": "Stress Test",
        "objective": "Verify robustness, availability, and error handling under adverse, boundary, or beyond-normal operating conditions.",
        "goals": "Detect crashes, unsafe states, unhandled exceptions, timing breakdowns, and poor recovery under stress.",
        "coverage": "Coverage includes boundary values, invalid inputs, timeouts, repeated faults, high load, and resource-limited behavior as applicable.",
        "technique": "Use controlled stress scenarios, negative tests, fault injection, abnormal timing, and recovery checks.",
        "responsible": "QA test engineer executes stress testing with safety and bench support where needed.",
        "special": "Run intrusive or safety-sensitive stress cases only in an approved controlled environment.",
        "steps": "Prepare safe stress conditions, apply adverse inputs/workloads, monitor safety and outputs, verify recovery, and log evidence.",
        "automation": "Automate controlled stress and fault-injection patterns where the bench supports repeatable execution.",
    },
    21: {
        "name": "Failure and Recovery Testing",
        "objective": "Verify that the software detects failures, enters the specified reaction, preserves safety, and recovers or reports status correctly.",
        "goals": "Detect recovery defects, unsafe transitions, incomplete diagnostics, retained-data issues, and poor behavior after interruptions.",
        "coverage": "Coverage includes fault-injection, interruption, timeout, diagnostic, reset/restart, safe-state, and recovery scenarios.",
        "technique": "Inject controlled faults or interruptions, monitor reactions and diagnostics, restore conditions, and verify return to a known acceptable state.",
        "responsible": "QA test engineer executes with systems, safety, and bench support.",
        "special": "Failure and recovery tests can be intrusive; isolate the bench and define rollback/safety controls before execution.",
        "steps": "Baseline normal state, inject failure, observe reaction, restore conditions, verify recovery/status, and archive logs/evidence.",
        "automation": "Automate safe fault-injection and recovery checks where controlled tooling is available.",
    },
    22: {
        "name": "Configuration/Compatibility Testing",
        "objective": "Verify operation on supported hardware, software, tool, interface, and configuration baselines.",
        "goals": "Detect compatibility defects, unsupported combinations, version mismatches, configuration gaps, and environment-specific failures.",
        "coverage": "Coverage includes approved hardware, software, communication artifacts, tools, calibration/configuration versions, and representative variants.",
        "technique": "Execute representative functional and smoke cases across selected supported configurations and record exact baselines.",
        "responsible": "QA test engineer executes with configuration-management and tool support.",
        "special": "Document all target and non-target software/tool versions used during compatibility testing.",
        "steps": "Select supported combinations, configure the environment, execute representative cases, compare results, and document compatibility status.",
        "automation": "Automate environment/version capture and repeatable compatibility smoke tests where practical.",
    },
    23: {
        "name": "Regression Testing",
        "objective": "Verify that unchanged and impacted behavior continues to pass after software or requirement changes.",
        "goals": "Detect reintroduced defects, side effects from fixes, interface regressions, and coverage gaps after changes.",
        "coverage": "Coverage includes sanity, changed requirements, impacted modules/interfaces, high-priority defects, and critical regression suites.",
        "technique": "Run selected and full regression suites based on change impact, priority, defect history, and release risk.",
        "responsible": "QA test engineer owns regression execution and reporting; developers support root-cause analysis.",
        "special": "Keep regression selection traceable to change impact and risk; rerun failed fixes before closure.",
        "steps": "Review change impact, select regression scope, execute automated/manual suites, triage failures, update RTM, and publish status.",
        "automation": "Automate stable sanity, interface, functional, and defect-prevention cases in the regression pipeline.",
    },
}


def _normalize_payload(payload: dict[str, Any]) -> dict[str, Any]:
    if "sections" in payload and isinstance(payload["sections"], dict):
        return payload["sections"]
    return payload


def _line_items(value: Any) -> list[str]:
    if isinstance(value, list):
        lines = []
        for item in value:
            if isinstance(item, dict):
                item_id = item.get("id", "")
                title = item.get("title", "")
                priority = item.get("priority", "")
                expected = item.get("expected", "")
                requirement_ids = ", ".join(item.get("requirementIds") or [])
                category = item.get("category", "")
                test_type = item.get("testType", "")
                preconditions = "; ".join(item.get("preconditions") or [])
                test_data = "; ".join(item.get("testData") or [])
                steps = "; ".join(item.get("steps") or [])
                prefix = f"{item_id} - " if item_id else ""
                suffix = f" ({priority})" if priority else ""
                details = [
                    f"Requirements: {requirement_ids}" if requirement_ids else "",
                    f"Category: {category}" if category else "",
                    f"Type: {test_type}" if test_type else "",
                    f"Preconditions: {preconditions}" if preconditions else "",
                    f"Test Data: {test_data}" if test_data else "",
                    f"Steps: {steps}" if steps else "",
                    f"Expected: {expected}" if expected else "",
                ]
                detail_text = " | ".join(detail for detail in details if detail)
                lines.append(f"{prefix}{title}{suffix}: {detail_text}".strip(": "))
            else:
                lines.append(str(item))
        return lines
    return [str(value)]


def render_pdf(payload: dict[str, Any]) -> bytes:
    sections = _normalize_payload(payload)
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=0.65 * inch,
        leftMargin=0.65 * inch,
        topMargin=0.65 * inch,
        bottomMargin=0.65 * inch,
        title="AI Generated Test Plan",
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "DocumentTitle",
        parent=styles["Title"],
        fontSize=20,
        leading=24,
        spaceAfter=18,
    )
    section_style = ParagraphStyle(
        "SectionTitle",
        parent=styles["Heading2"],
        fontSize=13,
        leading=16,
        spaceBefore=12,
        spaceAfter=8,
    )
    body_style = ParagraphStyle(
        "Body",
        parent=styles["BodyText"],
        fontSize=9.5,
        leading=13,
        spaceAfter=5,
    )

    story = [PdfParagraph("AI Generated Software Test Plan", title_style)]
    for key, title in SECTION_TITLES.items():
        value = sections.get(key)
        if not value:
            continue
        story.append(PdfParagraph(title, section_style))
        for line in _line_items(value):
            story.append(PdfParagraph(line.replace("\n", "<br/>"), body_style))
        story.append(Spacer(1, 6))

    doc.build(story)
    return buffer.getvalue()


def render_docx(payload: dict[str, Any]) -> bytes:
    sections = _normalize_payload(payload)
    document = Document(TEST_PLAN_TEMPLATE_PATH)
    _replace_template_placeholders(document, payload, sections)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _replace_template_placeholders(
    document: DocxDocument,
    payload: dict[str, Any],
    sections: dict[str, Any],
) -> None:
    heading_stack: list[str] = []
    table_index = 0

    _replace_paragraph_sequence_placeholders(document.paragraphs, payload, sections)

    for block in _iter_document_blocks(document):
        if isinstance(block, DocxParagraph):
            heading_context = " > ".join(heading_stack)
            _replace_paragraph_placeholders(
                block,
                payload,
                sections,
                heading_context=heading_context,
            )
            heading_level = _heading_level(block.style.name if block.style else "")
            if heading_level and block.text.strip():
                heading_stack = heading_stack[: heading_level - 1]
                heading_stack.append(_clean_text(block.text))
            continue

        table_index += 1
        heading_context = " > ".join(heading_stack)
        _replace_table_placeholders(
            block,
            payload,
            sections,
            table_index=table_index,
            heading_context=heading_context,
        )

    for section in document.sections:
        for header_footer in (
            section.header,
            section.first_page_header,
            section.even_page_header,
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        ):
            _replace_story_part_placeholders(header_footer, payload, sections)


def _iter_document_blocks(document: DocxDocument):
    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield DocxParagraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def _heading_level(style_name: str) -> int | None:
    if not style_name.startswith("Heading "):
        return None

    try:
        return int(style_name.removeprefix("Heading ").strip())
    except ValueError:
        return None


def _replace_story_part_placeholders(
    story_part: Any,
    payload: dict[str, Any],
    sections: dict[str, Any],
) -> None:
    _replace_paragraph_sequence_placeholders(story_part.paragraphs, payload, sections)

    for table in story_part.tables:
        _replace_table_placeholders(table, payload, sections)


def _replace_table_placeholders(
    table: Table,
    payload: dict[str, Any],
    sections: dict[str, Any],
    table_index: int | None = None,
    heading_context: str = "",
) -> None:
    test_context = TEST_APPROACH_CONTEXTS.get(table_index or 0)

    for row in table.rows:
        row_label = _clean_text(row.cells[0].text) if row.cells else ""
        for cell in row.cells:
            _replace_paragraph_sequence_placeholders(
                cell.paragraphs,
                payload,
                sections,
                heading_context=heading_context,
                row_label=row_label,
                test_context=test_context,
            )
            for nested_table in cell.tables:
                _replace_table_placeholders(
                    nested_table,
                    payload,
                    sections,
                    table_index=table_index,
                    heading_context=heading_context,
                )


def _replace_paragraph_placeholders(
    paragraph: DocxParagraph,
    payload: dict[str, Any],
    sections: dict[str, Any],
    heading_context: str = "",
    row_label: str = "",
    test_context: dict[str, str] | None = None,
) -> None:
    original_text = paragraph.text
    if "<" not in original_text or ">" not in original_text:
        return

    replaced_text = _replace_angle_placeholders(
        original_text,
        lambda description: _placeholder_replacement(
            description,
            payload,
            sections,
            heading_context=heading_context,
            row_label=row_label,
            test_context=test_context,
        ),
    )
    if replaced_text != original_text:
        _set_paragraph_text(paragraph, replaced_text)


def _replace_paragraph_sequence_placeholders(
    paragraphs: list[DocxParagraph],
    payload: dict[str, Any],
    sections: dict[str, Any],
    heading_context: str = "",
    row_label: str = "",
    test_context: dict[str, str] | None = None,
) -> None:
    index = 0
    while index < len(paragraphs):
        paragraph = paragraphs[index]
        text = paragraph.text

        if "<" in text and ">" not in text:
            end_index = index
            combined_parts = [text]
            found_close = False
            while end_index + 1 < len(paragraphs):
                next_text = paragraphs[end_index + 1].text
                if "<" in next_text:
                    break
                end_index += 1
                combined_parts.append(next_text)
                if ">" in next_text:
                    found_close = True
                    break

            combined_text = "\n".join(combined_parts)
            if "<" in combined_text and (found_close or ">" in combined_text):
                replaced_text = _replace_angle_placeholders(
                    combined_text,
                    lambda description: _placeholder_replacement(
                        description,
                        payload,
                        sections,
                        heading_context=heading_context,
                        row_label=row_label,
                        test_context=test_context,
                    ),
                )
                _set_paragraph_text(paragraph, replaced_text)
                for consumed_paragraph in paragraphs[index + 1 : end_index + 1]:
                    _set_paragraph_text(consumed_paragraph, "")
                index = end_index + 1
                continue
            if "<" in combined_text:
                start = combined_text.find("<")
                replaced_text = combined_text[:start] + _placeholder_replacement(
                    combined_text[start + 1 :],
                    payload,
                    sections,
                    heading_context=heading_context,
                    row_label=row_label,
                    test_context=test_context,
                )
                _set_paragraph_text(paragraph, replaced_text)
                for consumed_paragraph in paragraphs[index + 1 : end_index + 1]:
                    _set_paragraph_text(consumed_paragraph, "")
                index = end_index + 1
                continue

        _replace_paragraph_placeholders(
            paragraph,
            payload,
            sections,
            heading_context=heading_context,
            row_label=row_label,
            test_context=test_context,
        )
        index += 1


def _replace_angle_placeholders(text: str, resolve: Any) -> str:
    pieces: list[str] = []
    index = 0

    while index < len(text):
        start = text.find("<", index)
        if start == -1:
            pieces.append(text[index:])
            break

        opening_end = start
        while opening_end < len(text) and text[opening_end] == "<":
            opening_end += 1

        next_start = text.find("<", opening_end)
        segment_end = next_start if next_start != -1 else len(text)
        close_end = text.rfind(">", opening_end, segment_end)

        if close_end == -1:
            pieces.append(text[index:])
            break

        closing_start = close_end
        while closing_start > opening_end and text[closing_start - 1] == ">":
            closing_start -= 1

        pieces.append(text[index:start])
        pieces.append(resolve(text[opening_end:closing_start]))
        index = close_end + 1

    return "".join(pieces)


def _set_paragraph_text(paragraph: DocxParagraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
        return

    paragraph.add_run(text)


def _placeholder_replacement(
    description: str,
    payload: dict[str, Any],
    sections: dict[str, Any],
    heading_context: str = "",
    row_label: str = "",
    test_context: dict[str, str] | None = None,
) -> str:
    desc = _clean_placeholder_description(description)
    desc_l = desc.lower()
    heading_l = heading_context.lower()
    row_l = row_label.lower()

    if test_context:
        return _sanitize_replacement(
            _test_approach_replacement(desc_l, row_l, test_context, sections)
        )

    replacement = _general_placeholder_replacement(
        desc_l,
        heading_l,
        row_l,
        payload,
        sections,
    )
    return _sanitize_replacement(replacement)


def _test_approach_replacement(
    desc_l: str,
    row_l: str,
    context: dict[str, str],
    sections: dict[str, Any],
) -> str:
    if "responsible" in row_l or "who will perform" in desc_l:
        return context["responsible"]
    if "coverage" in row_l or "coverage" in desc_l:
        return context["coverage"]
    if "goal" in row_l or "goal" in desc_l:
        return context["goals"]
    if "technique" in row_l or "technique" in desc_l or "manual" in desc_l:
        return context["technique"]
    if "entry" in row_l or "entry criteria" in desc_l:
        return _format_items(
            _section_list(sections, "entryCriteria"),
            "Approved requirements, test cases, build, tools, test data, and environment are available before execution.",
            limit=4,
        )
    if "exit" in row_l or "exit criteria" in desc_l or "successfully completed" in desc_l:
        return _format_items(
            _section_list(sections, "exitCriteria"),
            "Planned tests pass or are formally deferred, critical defects are closed, and evidence is reviewed.",
            limit=4,
        )
    if "output" in row_l or "documentation" in desc_l:
        return _format_items(
            _section_list(sections, "deliverables"),
            "Test cases, execution logs, defect records, RTM updates, and test summary evidence.",
            limit=4,
        )
    if "setup" in row_l or "test setup" in desc_l:
        return _format_items(
            _section_list(sections, "testEnvironment"),
            "Approved project test environment with required hardware, software, tools, data, and logging.",
            limit=4,
        )
    if "automation" in row_l or "automation" in desc_l:
        return context["automation"]
    if "step" in row_l or "high level steps" in desc_l:
        return context["steps"]
    if "special" in row_l or "consideration" in desc_l:
        return context["special"]
    if "remark" in row_l or "additional information" in desc_l:
        return _phase_remarks(context, sections)

    return context["objective"]


def _general_placeholder_replacement(
    desc_l: str,
    heading_l: str,
    row_l: str,
    payload: dict[str, Any],
    sections: dict[str, Any],
) -> str:
    project_name = _project_name(payload, sections)

    if "project name" in desc_l:
        return project_name
    if "document history" in desc_l:
        return _document_history_text(payload, project_name)
    if "purpose of the test plan document" in desc_l:
        return _document_purpose(project_name, sections)
    if "scope of testing" in desc_l:
        return _section_text(
            sections,
            "scope",
            f"This test plan covers verification of {project_name} against the uploaded requirements.",
        )
    if "software overview" in desc_l:
        return _software_overview(project_name, sections)
    if "opportunity plan" in desc_l or "svn path" in desc_l:
        return (
            "PMR/Opportunity Plan repository path is not specified in the source document; "
            "confirm the controlled SVN/project repository location with the project manager."
        )
    if "known issues" in desc_l:
        return _known_issues(sections)
    if "thought out the approach" in desc_l or "test strategy" in heading_l:
        return _section_text(sections, "testStrategy", _default_strategy(project_name))
    if desc_l == "provide the path":
        return "Store the test plan, RTM, test cases, evidence, and related review artifacts in the approved project repository path."
    if "testing deliverables in each phase" in desc_l:
        return _phase_deliverables(sections)
    if "additional types of testing" in desc_l:
        return _testing_types_summary(sections)
    if "different types of testing" in desc_l:
        return _testing_types_summary(sections)
    if "mind map" in desc_l:
        return (
            "Create or update the product and feature understanding mind map in the project repository "
            "and link it to requirement-level test scenarios and the RTM."
        )
    if "mandatory for automotive" in desc_l and "integration is within" in desc_l:
        return _integration_applicability(sections)
    if "mandatory for automotive" in desc_l:
        return _automotive_applicability(sections)
    if "integration strategy" in desc_l:
        return _integration_strategy_text(sections)
    if "writing down the test case" in desc_l or "naming convention" in desc_l:
        return _test_case_naming_text(sections)
    if "source code" in desc_l or "automation script verification" in desc_l:
        return TEST_APPROACH_CONTEXTS[9]["objective"]
    if "features to be tested" in desc_l or "list all the features" in desc_l:
        return _format_items(
            _section_list(sections, "featuresToTest"),
            "Requirements and features identified from the uploaded source document.",
            limit=12,
        )
    if "features not" in desc_l or "list of features not" in desc_l:
        return _format_items(
            _section_list(sections, "featuresNotToTest"),
            "No explicit exclusions were provided; confirm any out-of-scope items with stakeholders.",
            limit=8,
        )
    if "priority for testing" in desc_l:
        return _priority_text(sections)
    if "sanity test" in desc_l:
        return TEST_APPROACH_CONTEXTS[15]["objective"] + " " + TEST_APPROACH_CONTEXTS[15]["coverage"]
    if "installation testing" in desc_l:
        return TEST_APPROACH_CONTEXTS[16]["objective"] + " " + TEST_APPROACH_CONTEXTS[16]["special"]
    if "performance profiling" in desc_l or "performance testing" in heading_l:
        return _section_text(
            sections,
            "nonFunctionalTesting",
            TEST_APPROACH_CONTEXTS[18]["objective"],
        )
    if "load testing" in desc_l:
        return TEST_APPROACH_CONTEXTS[19]["objective"] + " " + TEST_APPROACH_CONTEXTS[19]["coverage"]
    if "stress test" in desc_l:
        return TEST_APPROACH_CONTEXTS[20]["objective"] + " " + TEST_APPROACH_CONTEXTS[20]["coverage"]
    if "configuration testing" in desc_l:
        return TEST_APPROACH_CONTEXTS[22]["objective"] + " " + TEST_APPROACH_CONTEXTS[22]["coverage"]
    if "test environment" in desc_l or "test setup" in desc_l:
        return _format_items(
            _section_list(sections, "testEnvironment"),
            "Approved project test environment with required hardware, software, tools, data, and logging.",
            limit=6,
        )
    if "test setup diagrams" in desc_l or "setup diagram" in desc_l:
        return (
            "Use the approved SIL/HIL/bench setup diagram for the project. If the source document does "
            "not provide a diagram, add a diagram showing the software under test, simulated/actual "
            "external interfaces, diagnostic tools, test controller, logging, and power or communication dependencies."
        )
    if "remarks column" in desc_l:
        return "Record the applicable phase, baseline/version, access constraints, and setup remarks for each item."
    if "hardware dependency" in desc_l or row_l == "hardware":
        return _hardware_dependency(sections)
    if "software dependency" in desc_l or row_l == "software":
        return _software_dependency(sections)
    if "test data" in desc_l or row_l == "test data":
        return _test_data_dependency(sections)
    if "test tool" in desc_l or "testing tools" in row_l:
        return _tool_dependency(sections)
    if "pass fail" in desc_l or "pass/fail" in heading_l:
        return _pass_fail_criteria(sections)
    if "suspended" in desc_l:
        return _suspension_criteria(sections)
    if "decided by the tester" in desc_l or "resumption" in heading_l:
        return _resumption_criteria(sections)
    if "as provide by the client" in desc_l or "acceptance criteria" in heading_l:
        return _acceptance_criteria(sections)
    if "work products selected for validation" in desc_l:
        return _format_items(
            _validation_work_products(sections),
            "SRS, RTM, test plan, test cases, execution evidence, defect reports, and test summary.",
            limit=8,
        )
    if "validation methods" in desc_l:
        return _format_items(
            [
                "Review and baseline of requirements, test plan, and RTM.",
                "Requirement-based test execution with objective evidence.",
                "Defect verification and regression confirmation.",
                "Stakeholder review and approval of validation results.",
            ],
            "",
        )
    if "validation environment" in desc_l:
        return _format_items(
            _section_list(sections, "testEnvironment"),
            "Approved project validation environment representing the intended SIL/HIL/bench or target setup.",
            limit=5,
        )
    if "validation criteria" in desc_l:
        return _format_items(
            [
                "Requirements are traceable to test cases and execution status.",
                "High-priority and critical tests pass with reviewed evidence.",
                "Open defects are dispositioned and accepted by stakeholders.",
                "Release recommendation is supported by RTM and test summary results.",
            ],
            "",
        )
    if "remove rows where not applicable" in desc_l:
        return (
            "Retain rows for requirement, design, coding, testing, release, and validation work products "
            "that are in project scope; mark non-applicable rows as not applicable or remove them during document tailoring."
        )
    if "bug" in desc_l and "severity" in desc_l:
        return (
            "Defects shall record summary, detailed description, requirement/test case reference, severity, "
            "priority, environment, build version, evidence, reproduction steps, owner, status, and closure notes."
        )
    if "summary field" in desc_l:
        return "Use a concise one-line summary describing the observed failure, affected feature, and trigger condition."
    if "additional information about the bug" in desc_l:
        return "Include logs, screenshots/captures, build and tool versions, environment details, similar prior defects, and retest notes."
    if "status reporting" in desc_l:
        return (
            "Generate status from the defect tracking system, RTM, and test case execution records. "
            "Report planned/executed/passed/failed/blocked counts, open defect severity, risks, and release recommendation."
        )
    if "training requirement" in desc_l:
        return _training_requirements(sections)
    if desc_l == "if any":
        return "No additional appendix items are identified from the source document; add project-specific supporting material if required."

    return _inferred_general_text(desc_l, project_name, sections)


def _clean_placeholder_description(value: str) -> str:
    text = value.replace("\xa0", " ").replace("<", " ").replace(">", " ")
    text = text.replace("\n", " / ")
    return _clean_text(text.strip(" /"))


def _sanitize_replacement(value: str) -> str:
    text = str(value or "").replace("\xa0", " ").replace("<", "").replace(">", "")
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in text.splitlines()]
    text = "\n".join(line for line in lines if line).strip()
    return text or "Not specified in the source document; confirm with stakeholders."


def _clean_text(value: Any) -> str:
    if value is None:
        return ""
    return re.sub(r"\s+", " ", str(value)).strip()


def _source_files(payload: dict[str, Any]) -> list[str]:
    source_files = payload.get("sourceFiles")
    if not isinstance(source_files, list):
        return []
    return [_clean_text(file_name) for file_name in source_files if _clean_text(file_name)]


def _project_name(payload: dict[str, Any], sections: dict[str, Any]) -> str:
    generated_text = " ".join(
        [
            _section_text(sections, "scope", ""),
            _section_text(sections, "testStrategy", ""),
            " ".join(_section_list(sections, "featuresToTest")[:5]),
        ]
    )
    patterns = (
        r"\bfor the ([A-Z][A-Za-z0-9 /_-]{2,80}? software)\b",
        r"\bfor ([A-Z][A-Za-z0-9 /_-]{2,80}? software)\b",
        r"\b([A-Z0-9]{2,}(?: [A-Z0-9]{2,}){1,5}) software\b",
    )
    for pattern in patterns:
        match = re.search(pattern, generated_text)
        if match:
            return _title_preserving_acronyms(match.group(1))

    source_files = _source_files(payload)
    if source_files:
        stem = Path(source_files[0]).stem.replace("_", " ").replace("-", " ")
        cleaned = _clean_text(re.sub(r"\b(srs|swrs|requirements?|document|actual|example)\b", "", stem, flags=re.I))
        if cleaned:
            return _title_preserving_acronyms(cleaned)

    return "Software Test Plan"


def _title_preserving_acronyms(value: str) -> str:
    words = []
    for word in _clean_text(value).split():
        if len(word) > 1 and any(character.isdigit() for character in word):
            words.append(word.upper())
        elif word.isupper() and len(word) > 1:
            words.append(word)
        else:
            words.append(word[:1].upper() + word[1:])
    return " ".join(words)


def _document_history_text(payload: dict[str, Any], project_name: str) -> str:
    date_text = _created_date(payload)
    source_files = ", ".join(_source_files(payload)) or "uploaded source document"
    return (
        f"Initial generated test plan baseline for {project_name}. Generated on {date_text} "
        f"from {source_files}; update version, author, reviewer, and approval fields per project governance."
    )


def _document_purpose(project_name: str, sections: dict[str, Any]) -> str:
    scope = _section_text(
        sections,
        "scope",
        f"verification of {project_name} against the uploaded requirements",
    )
    return (
        f"The purpose of this test plan is to define the project-specific testing approach for {project_name}. "
        f"It identifies the test scope, strategy, test levels, environment, entry and exit criteria, risks, "
        f"deliverables, traceability expectations, and generated test scenarios/test cases needed to validate {scope}"
    )


def _phase_deliverables(sections: dict[str, Any]) -> str:
    common_deliverables = _section_list(sections, "deliverables")
    common_text = _format_inline(
        common_deliverables,
        "RTM, detailed test cases, execution evidence, defect records, and final test summary",
        limit=4,
    )
    return _format_items(
        [
            "Planning and analysis: software test plan, reviewed scope, risks, assumptions, dependencies, and RTM baseline.",
            "Test design: detailed requirement-level test cases with preconditions, data, steps, expected results, priority, and traceability.",
            "Test execution: executed test records, logs/captures, defect reports, retest evidence, and updated RTM status.",
            f"Closure: final test summary, open risk/deviation list, release recommendation, and supporting deliverables such as {common_text}",
        ],
        "",
    )


def _created_date(payload: dict[str, Any]) -> str:
    created_at = _clean_text(payload.get("createdAt"))
    if not created_at:
        return "the export date"

    try:
        return created_at.split("T", 1)[0]
    except (AttributeError, IndexError):
        return "the export date"


def _section_text(sections: dict[str, Any], key: str, fallback: str) -> str:
    value = sections.get(key)
    if isinstance(value, list):
        text = " ".join(_line_items(value))
    else:
        text = _clean_text(value)
    return text or fallback


def _section_list(sections: dict[str, Any], key: str) -> list[str]:
    value = sections.get(key)
    if isinstance(value, list):
        return [_clean_text(item) for item in _line_items(value) if _clean_text(item)]
    if value:
        return [_clean_text(value)]
    return []


def _format_items(items: list[str], fallback: str, limit: int | None = None) -> str:
    cleaned = [_clean_text(item) for item in items if _clean_text(item)]
    if limit is not None:
        cleaned = cleaned[:limit]
    if not cleaned:
        return fallback
    return "\n".join(f"- {item}" for item in cleaned)


def _format_inline(items: list[str], fallback: str, limit: int = 4) -> str:
    cleaned = [_clean_text(item).rstrip(".") for item in items if _clean_text(item)]
    if not cleaned:
        return fallback
    return "; ".join(cleaned[:limit]) + "."


def _software_overview(project_name: str, sections: dict[str, Any]) -> str:
    features = _format_inline(
        _section_list(sections, "featuresToTest"),
        "the requirements and features extracted from the uploaded source document",
        limit=5,
    )
    interfaces = _format_inline(
        _section_list(sections, "apiTesting") + _section_list(sections, "testEnvironment"),
        "the approved project interfaces and verification environment",
        limit=3,
    )
    return (
        f"{project_name} is the software under test for this plan. The generated scope focuses on {features} "
        f"High-level verification includes {interfaces} Add or reference the project block diagram showing "
        "the software under test, external interfaces, diagnostic/test tools, data sources, and SIL/HIL/bench environment."
    )


def _known_issues(sections: dict[str, Any]) -> str:
    risks = _section_list(sections, "risks")
    if risks:
        return (
            "No source-specific open issues are explicitly listed. Track defects in the project defect system. "
            "Initial risk-driven watch items are:\n" + _format_items(risks, "", limit=5)
        )
    return "No known issues are explicitly identified in the source document; track newly found issues in the defect system."


def _default_strategy(project_name: str) -> str:
    return (
        f"Use a requirements-based test strategy for {project_name}, maintaining traceability from source "
        "requirements to test cases, execution evidence, defects, and release recommendation."
    )


def _testing_types_summary(sections: dict[str, Any]) -> str:
    type_lines = [
        TEST_APPROACH_CONTEXTS[10]["name"] + ": " + TEST_APPROACH_CONTEXTS[10]["objective"],
        TEST_APPROACH_CONTEXTS[12]["name"] + ": " + TEST_APPROACH_CONTEXTS[12]["objective"],
        TEST_APPROACH_CONTEXTS[14]["name"] + ": " + TEST_APPROACH_CONTEXTS[14]["objective"],
        TEST_APPROACH_CONTEXTS[17]["name"] + ": " + _section_text(sections, "functionalTesting", TEST_APPROACH_CONTEXTS[17]["objective"]),
        TEST_APPROACH_CONTEXTS[18]["name"] + ": " + _section_text(sections, "nonFunctionalTesting", TEST_APPROACH_CONTEXTS[18]["objective"]),
        TEST_APPROACH_CONTEXTS[23]["name"] + ": " + _section_text(sections, "regressionTesting", TEST_APPROACH_CONTEXTS[23]["objective"]),
    ]
    if _section_list(sections, "securityTesting"):
        type_lines.append("Security/robustness testing: " + _section_text(sections, "securityTesting", ""))
    if _section_list(sections, "apiTesting"):
        type_lines.append("Interface/API testing: " + _section_text(sections, "apiTesting", ""))
    return _format_items(type_lines, "")


def _automotive_applicability(sections: dict[str, Any]) -> str:
    if _is_automotive_context(sections):
        return (
            "Applicable. The project context indicates automotive or embedded software verification; "
            "execute this section with requirement traceability, safety/diagnostic evidence, configuration control, and regression coverage."
        )
    return (
        "Applicability is not explicit in the source document. Confirm automotive standard applicability "
        "with stakeholders and execute this section when required by the project scope."
    )


def _integration_applicability(sections: dict[str, Any]) -> str:
    if _section_list(sections, "apiTesting") or "interface" in _all_sections_text(sections):
        return (
            "Applicable where interface or component integration is in scope. Verify integrated software "
            "behavior, external interfaces, diagnostics, configuration compatibility, and regression impacts."
        )
    return "Integration scope is not explicit in the source document; confirm applicability with stakeholders."


def _integration_strategy_text(sections: dict[str, Any]) -> str:
    interface_scope = _format_inline(
        _section_list(sections, "apiTesting") + _section_list(sections, "featuresToTest"),
        "the interfaces, components, and features identified from the uploaded requirements",
        limit=5,
    )
    return (
        "Use an incremental, requirement-driven integration strategy. Integrate components and interfaces "
        f"according to build availability and risk, with priority on {interface_scope} Manage the integration "
        "environment through controlled software, configuration, interface-artifact, tool, and test-data baselines. "
        "Evaluate each build with smoke, interface, diagnostic, functional, negative, and regression checks, and "
        "record results in the RTM with linked defects and evidence."
    )


def _test_case_naming_text(sections: dict[str, Any]) -> str:
    categories = _section_list(sections, "featuresToTest")
    category_hint = _format_inline(categories, "the affected feature or interface", limit=3)
    return (
        "Use a consistent requirement and interface-oriented naming convention for integration tests. "
        "Recommended format: IT_<feature_or_interface>_<requirement_id>_<sequence>, where the feature/interface "
        f"name is based on {category_hint} Keep generated test case IDs traceable to the RTM and source requirement IDs."
    )


def _is_automotive_context(sections: dict[str, Any]) -> bool:
    text = _all_sections_text(sections)
    indicators = ("automotive", "vehicle", "can", "dbc", "arxml", "obd", "hil", "sil", "iso 26262", "hv", "dcdc")
    return any(indicator in text for indicator in indicators)


def _all_sections_text(sections: dict[str, Any]) -> str:
    return " ".join(_section_text(sections, key, "") for key in SECTION_TITLES).lower()


def _priority_text(sections: dict[str, Any]) -> str:
    test_cases = sections.get("testCases") if isinstance(sections.get("testCases"), list) else []
    high_count = sum(1 for item in test_cases if isinstance(item, dict) and _clean_text(item.get("priority")).lower() == "high")
    return (
        "Prioritize testing by requirement criticality, safety impact, customer acceptance criteria, changed/impacted areas, "
        f"and defect history. Execute sanity checks first, then high-priority cases ({high_count} currently generated), "
        "previous defect retests, requirement-based functional/interface cases, non-functional cases, and exploratory testing."
    )


def _hardware_dependency(sections: dict[str, Any]) -> str:
    return _format_inline(
        _section_list(sections, "testEnvironment"),
        "Required target hardware, bench/HIL/SIL setup, communication interfaces, power supplies, and logging equipment must be available.",
        limit=3,
    )


def _software_dependency(sections: dict[str, Any]) -> str:
    return (
        "Approved software build, configuration/calibration baseline, interface artifacts, diagnostic data, "
        "test scripts, and source requirement baseline must be available before execution."
    )


def _test_data_dependency(sections: dict[str, Any]) -> str:
    cases = sections.get("testCases") if isinstance(sections.get("testCases"), list) else []
    data_items: list[str] = []
    for case in cases:
        if isinstance(case, dict):
            data_items.extend(_clean_text(item) for item in case.get("testData", []) if _clean_text(item))
    return _format_inline(
        data_items,
        "Use SRS-specified signals, states, diagnostic requests, calibration values, boundary values, and expected results.",
        limit=5,
    )


def _tool_dependency(sections: dict[str, Any]) -> str:
    return _format_inline(
        _section_list(sections, "testEnvironment"),
        "Required test management, diagnostic, communication simulation, automation, static-analysis, logging, and defect tracking tools must be available.",
        limit=4,
    )


def _pass_fail_criteria(sections: dict[str, Any]) -> str:
    return (
        "A test passes when actual behavior matches the expected result, acceptance criteria, logs, timing, "
        "diagnostics, and interface outputs defined for the requirement. A test fails when any mandatory "
        "expected result is not met, evidence is incomplete, or an unexpected defect/unsafe behavior is observed. "
        + _format_inline(_section_list(sections, "exitCriteria"), "", limit=3)
    ).strip()


def _suspension_criteria(sections: dict[str, Any]) -> str:
    return (
        "Suspend affected testing when blocker or critical defects prevent meaningful execution, the environment "
        "is unstable, required hardware/software/tools are unavailable, safety controls are not satisfied, or source "
        "requirements/configuration baselines are inconsistent. "
        + _format_inline(_section_list(sections, "risks"), "", limit=2)
    ).strip()


def _resumption_criteria(sections: dict[str, Any]) -> str:
    return (
        "Resume testing after blocker defects are fixed or formally dispositioned, the environment is restored, "
        "required baselines and data are available, and impacted sanity/regression checks pass."
    )


def _acceptance_criteria(sections: dict[str, Any]) -> str:
    return _format_items(
        _section_list(sections, "exitCriteria"),
        "Acceptance requires completion of planned tests, closure or approved deferral of critical defects, reviewed evidence, and stakeholder approval.",
        limit=6,
    )


def _validation_work_products(sections: dict[str, Any]) -> list[str]:
    work_products = [
        "Source requirements/SRS baseline",
        "Requirement traceability matrix",
        "Software test plan",
        "Detailed test cases",
        "Test execution evidence and logs",
        "Defect reports",
        "Final test summary report",
    ]
    if _section_list(sections, "deliverables"):
        work_products.extend(_section_list(sections, "deliverables")[:4])
    return work_products


def _training_requirements(sections: dict[str, Any]) -> str:
    focus = _format_inline(
        _section_list(sections, "testEnvironment") + _section_list(sections, "apiTesting"),
        "project requirements, test environment, tools, domain behavior, defect workflow, and safety/security procedures",
        limit=4,
    )
    return (
        "Test engineers shall be trained on the source requirements, generated RTM/test cases, execution workflow, "
        f"and {focus}"
    )


def _phase_remarks(context: dict[str, str], sections: dict[str, Any]) -> str:
    remarks = [
        f"Execute {context['name']} using the approved baseline and retain objective evidence.",
        "Link results to requirement IDs, test cases, defects, and RTM status.",
    ]
    risks = _section_list(sections, "risks")
    if risks:
        remarks.append("Review risk controls before execution: " + "; ".join(risks[:2]) + ".")
    if _is_automotive_context(sections):
        remarks.append("Confirm safety-related checks do not violate defined safety goals or controlled bench limits.")
    return " ".join(remarks)


def _inferred_general_text(desc_l: str, project_name: str, sections: dict[str, Any]) -> str:
    if "entry" in desc_l:
        return _format_items(_section_list(sections, "entryCriteria"), "Approved inputs are available before execution.", limit=4)
    if "exit" in desc_l:
        return _format_items(_section_list(sections, "exitCriteria"), "Planned verification is complete and reviewed.", limit=4)
    if "documentation" in desc_l or "deliverable" in desc_l:
        return _format_items(_section_list(sections, "deliverables"), "Test plan, test cases, RTM, execution evidence, defect reports, and summary report.", limit=5)
    if "objective" in desc_l:
        return _default_strategy(project_name)
    if "technique" in desc_l:
        return "Use requirement-based positive, negative, boundary, interface, regression, and risk-based testing techniques as applicable."
    if "consideration" in desc_l:
        return "Maintain configuration control, traceability, safety controls, tool readiness, and reviewed evidence for this activity."

    return (
        "Information is not explicitly available in the source document; use the generated requirements, "
        "test strategy, RTM, and stakeholder-confirmed project baseline to complete this field consistently."
    )
